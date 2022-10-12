#

# collect frames (both syn and sem) and process data for foreign ones
# note: mainly for zh/es conll09

from collections import Counter, OrderedDict
from msp2.data.inst import yield_sents, set_ee_heads
from msp2.utils import Conf, zlog, zopen, zwarn, init_everything, default_json_serializer, OtherHelper, zglob
from msp2.data.rw import ReaderGetterConf, WriterGetterConf
from msp2.tasks.zmtl3.mod.extract.evt_arg import onto as zonto

class MainConf(Conf):
    def __init__(self):
        self.output_onto = ""
        self.input_file = ""  # input data
        self.language = 'UNK'
        self.ref_onto = ""
        self.tfile = ''  # translation file: each line "{src}\t{trg}"
        self.add_dummy_tpl = False
        self.rm_no_args = False
        # --
        self.role_thr = 0.05  # role appear larger than this
        # --

# --
def get_role_np(role: str, **extra_info):
    d = {
        'nsubj': 'subject', 'obj': 'object', 'iobj': 'indirect-object', 'obl': 'oblique',
        'nmod': 'modifier', 'compound': 'compound',
        'ARGM-LOC': 'place',
    }
    if role in d:
        return d[role]
    # --
    assert role.startswith("ARG") and str.isdigit(role[3])
    # --
    d2 = {
        'adv': 'adverb', 'agt': 'agent', 'atr': 'attribute', 'ben': 'beneficiary', 'cau': 'cause', 'cot': 'co-theme',
        'des': 'destination', 'efi': 'state', 'ein': 'state', 'exp': 'experiencer', 'ext': 'extension', 'fin': 'purpose',
        'ins': 'instrument', 'loc': 'place', 'mnr': 'manner', 'ori': 'origin', 'pat': 'patient', 'src': 'source',
        'tem': 'theme', 'tmp': 'time',
        'null': 'entity',  # special one
    }
    if role.split('-')[-1].lower() in d2:
        return d2[role.split('-')[-1].lower()]
    # --
    # finally UNK one?
    zwarn(f"UNK role of {role} with {extra_info}")
    # return "entity"  # dummy one
    return role  # return original!

# --
# helper function to export frame names
def export_frame_names(input_file: str, output_file: str):
    from msp2.data.rw import ReaderGetterConf
    from collections import Counter
    from msp2.data.inst import yield_frames
    from docx import Document
    names = Counter()
    for inst in ReaderGetterConf().get_reader(input_path=input_file):
        for frame in yield_frames(inst):
            one_name = frame.label.split(".")[0]
            names[one_name] += 1
    with open(output_file, 'w') as fd:
        for a, b in names.most_common():
            fd.write(a+"\n")
    # also write docx
    document = Document()
    for a, b in names.most_common():
        document.add_paragraph(a)
    document.save(output_file+".docx")
    # --
# --
def docx2plain(filenames, output_file):
    from docx import Document
    import stanza
    nlp = stanza.Pipeline(lang='en', processors='tokenize,mwt,pos,lemma', use_gpu=False)
    # --
    all_lines = None
    for f in filenames:
        do_nlp = "_en." in f
        print(f"Read {f}: {do_nlp}")
        lines = []
        doc = Document(f)
        for para in doc.paragraphs:
            one = para.text.strip().lower()
            # simple parse and norm
            if do_nlp and len(one.split()) > 1:
                _doc = nlp(one)  # simply check pos
                _words = [word for sent in _doc.sentences for word in sent.words]
                # skipping starting PRONs/TO and put lemma
                _ii = 0
                while _ii < len(_words) and (_words[_ii].upos == "PRON" or _words[_ii].text == 'to'):
                    _ii += 1
                one2 = " ".join([w.lemma for w in _words[_ii:] if w.lemma is not None])
                if len(one2) > 0:
                    one = one2
            # --
            lines.append(one)
        if all_lines is None:
            all_lines = lines
        else:
            assert len(lines) == len(all_lines)
            for ii in range(len(all_lines)):
                all_lines[ii] = all_lines[ii] + "\t" + lines[ii]
    with open(output_file, 'w') as fd:
        for line in all_lines:
            fd.write(line+"\n")
    # --
# --

# --
def main(*args):
    conf: MainConf = init_everything(MainConf(), args)
    # --
    # first read data
    reader = ReaderGetterConf().get_reader(input_path=conf.input_file)
    all_insts = list(reader)
    ref_onto = None
    if conf.ref_onto:
        ref_onto = zonto.Onto.load_onto(conf.ref_onto)
    tmap = None
    if conf.tfile:
        tmap = {}
        with zopen(conf.tfile) as fd:
            for line in fd:
                try:
                    a, b = line.strip().split("\t")
                except:
                    zwarn(f"Skip strange line: {line}")
                if a not in tmap:
                    tmap[a] = b
        zlog(f"Read tfile {conf.tfile}: {len(tmap)}")
    # --
    # read data
    frame_role_ccs = {}
    cc = Counter()
    for inst in all_insts:
        cc['inst'] += 1
        for sent in yield_sents(inst):
            cc['sent'] += 1
            for evt in sent.events:
                cc['evt'] += 1
                cc['arg'] += len(evt.args)
                # --
                evt_type = evt.label
                if evt_type not in frame_role_ccs:
                    frame_role_ccs[evt_type] = [0, Counter()]
                frame_role_ccs[evt_type][0] += 1
                frame_role_ccs[evt_type][1].update([z.label for z in evt.args])
    # --
    # put onto
    all_frames, all_roles = [], []
    for name in sorted(frame_role_ccs.keys(), key=(lambda x: frame_role_ccs[x][0]), reverse=True):
        cc['oframe'] += 1
        # --
        ref_frame = None
        if ref_onto is not None:
            ref_frame = ref_onto.find_frame(name)
            if ref_frame is None:
                cc['oframeRF'] += 1
                zwarn(f"Failed finding ref-frame of {name}")
                continue
        # --
        _ccs = frame_role_ccs[name]
        roles = []
        for _role, _count in _ccs[1].items():
            cc['oroleA'] += 1
            if _count / _ccs[0] >= conf.role_thr:
                cc['oroleI'] += 1
                # --
                _np = None
                if ref_frame is not None:  # try find it from ref
                    ref_role = ref_frame.find_role(_role, None)[0]
                    if ref_role is not None:
                        _np = ref_role.np
                if _np is None:
                    _np = get_role_np(_role, fname=name, _ccs=_ccs)
                rr = zonto.Role(_role, np=_np)
                roles.append(rr)
            else:
                cc['oroleE'] += 1  # exclude!
        # --
        _vp = name.split(".")[0]
        if tmap:
            _vp = tmap.get(_vp)
            if _vp is None:
                cc['oframeTF'] += 1
                zwarn(f"Failed translation of {name}/{_vp}")
                continue
        ff = zonto.Frame(name, vp=_vp, core_roles=roles)  # note: all core!
        # --
        # add a dummy template anyway
        if conf.add_dummy_tpl:
            _tmp_tpl = [[z, []] for z in sorted([rr.name for rr in roles])]
            ff.template = _tmp_tpl[:1] + [[None, []]] + _tmp_tpl[1:]
        # --
        if conf.rm_no_args and len(roles) == 0:
            cc['oframeNoA'] += 1
            continue
        all_frames.append(ff)
        all_roles.extend(roles)
    onto = zonto.Onto(all_frames, all_roles)
    zlog(f"Collect onto {onto} from {conf.input_file} to {conf.output_onto}: {cc}")
    # --
    # write
    if conf.output_onto:
        default_json_serializer.to_file(onto.to_json(), conf.output_onto, indent=2)
    # --

# python3 -m msp2.tasks.zmtl3.scripts.srl.sz_collect_onto ...
if __name__ == '__main__':
    import sys
    main(*sys.argv[1:])
